"""
Document Generator Service
Generates PDF and DOCX documents from application data
"""
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class DocumentGenerator:
    """Service for generating application documents"""
    
    def generate_docx(self, application, output_path: str) -> str:
        """
        Generate DOCX document from application
        
        Args:
            application: Application model instance
            output_path: Path where to save the document
            
        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading(application.project_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Antrag ID: {application.id}")
        doc.add_paragraph(f"Erstellt am: {application.created_at.strftime('%d.%m.%Y')}")
        doc.add_paragraph("")
        
        # Executive Summary
        doc.add_heading("Zusammenfassung", level=1)
        doc.add_paragraph(application.project_description)
        doc.add_paragraph("")
        
        # Generate sections if available
        if application.generated_content:
            sections = [
                ("Projektbeschreibung", "project_description"),
                ("Marktanalyse", "market_analysis"),
                ("Technische Machbarkeit", "technical_feasibility"),
                ("Arbeitsplan", "work_plan"),
                ("Finanzplan", "financial_plan"),
                ("Risikomanagement", "risk_management"),
                ("Verwertungsplan", "utilization_plan")
            ]
            
            for heading, key in sections:
                if key in application.generated_content:
                    doc.add_heading(heading, level=1)
                    content = application.generated_content[key]
                    doc.add_paragraph(content)
                    doc.add_paragraph("")
        
        # Budget overview
        doc.add_heading("Budget-Übersicht", level=1)
        budget_table = doc.add_table(rows=4, cols=2)
        budget_table.style = 'Light Grid Accent 1'
        
        budget_data = [
            ("Gesamtbudget", f"{application.total_budget:,.2f} €"),
            ("Beantragte Förderung", f"{application.requested_funding:,.2f} €"),
            ("Eigenmittel", f"{application.own_contribution:,.2f} €"),
            ("Laufzeit", f"{application.timeline_months} Monate")
        ]
        
        for i, (label, value) in enumerate(budget_data):
            budget_table.rows[i].cells[0].text = label
            budget_table.rows[i].cells[1].text = value
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def generate_pdf(self, application, output_path: str) -> str:
        """
        Generate PDF document from application
        
        For now, generates DOCX and converts to PDF (requires LibreOffice)
        In production, use proper PDF library like ReportLab or WeasyPrint
        
        Args:
            application: Application model instance
            output_path: Path where to save the PDF
            
        Returns:
            Path to generated PDF
        """
        # Generate DOCX first
        docx_path = output_path.replace('.pdf', '.docx')
        self.generate_docx(application, docx_path)
        
        # For MVP: Just return DOCX path
        # In production: Convert to PDF using LibreOffice or ReportLab
        # Example: subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_path])
        
        # For now, return DOCX (rename to indicate it's actually DOCX)
        return docx_path


# Singleton instance
document_generator = DocumentGenerator()
